"""
Назначение файла: разделы Word-заключения.

Что выполняет файл:
- добавляет в документ шапку, данные пациента, опрос, осмотр, анализы, диагнозы,
  назначения и подпись;
- сохраняет текущий порядок и текстовую структуру заключения;
- не загружает данные из БД и не отдаёт HTTP-ответ.
"""

from __future__ import annotations

from docx.shared import Pt

from .formatting import (
    add_centered_paragraph,
    add_field_inline,
    add_history_table,
    add_small_table,
    add_table_title,
    clean_value,
    fmt_date,
    clean_word_recommendations,
    icd10_diagnoses_conclusion_text,
    kdigo_conclusion_text,
    prognosis_display,
    set_run_font,
)


def add_clinic_header(doc, location_info):
    """Шапка клиники в две строки."""
    if location_info:
        company_name = location_info.get("company_name") or "ООО «КОМПАНИЯ «ФЕСФАРМ»"
        location_name = location_info.get("location_name") or ""
        location_address = location_info.get("location_address") or ""
        branch_phone = location_info.get("branch_phone") or ""
        company_phone = location_info.get("company_phone") or ""
        phone = branch_phone or company_phone or ""
        branch_email = location_info.get("branch_email") or ""
        company_email = location_info.get("company_email") or ""
        email = branch_email or company_email or ""

        first_line = company_name
        if location_name:
            first_line += f" — {location_name}"

        second_line_parts = []
        if location_address:
            second_line_parts.append(location_address)
        if phone:
            second_line_parts.append(f"Тел: {phone}")
        if email:
            second_line_parts.append(f"Email: {email}")

        add_centered_paragraph(doc, first_line, size=9, bold=False, space_after=0)
        add_centered_paragraph(doc, " | ".join(second_line_parts), size=9, bold=False, space_after=4)
    else:
        add_centered_paragraph(doc, "ООО «КОМПАНИЯ «ФЕСФАРМ»", size=9, bold=False, space_after=4)


def add_document_title(doc):
    add_centered_paragraph(doc, "КОНСУЛЬТАТИВНОЕ ЗАКЛЮЧЕНИЕ", size=14, bold=True, space_after=6)


def add_patient_section(doc, appointment):
    add_field_inline(doc, "Пациент", appointment.get("patient_fio"))
    add_field_inline(doc, "Дата рождения", fmt_date(appointment.get("birth_date")))

    if appointment.get("age_at_appointment") is not None:
        try:
            age_text = f"{int(appointment.get('age_at_appointment'))} лет"
        except Exception:
            age_text = str(appointment.get("age_at_appointment"))
        add_field_inline(doc, "Возраст на момент приёма", age_text)


def _sentence(value: str) -> str:
    """Нормализует фрагмент текста до одной фразы с точкой в конце."""
    value = (value or "").strip()
    if not value:
        return ""
    value = value.rstrip(" ;")
    if value[-1] not in ".!?":
        value += "."
    return value


def _join_sentences(parts) -> str:
    return " ".join(part for part in parts if part).strip()


def add_survey_section(doc, appointment):
    add_field_inline(doc, "Жалобы", appointment.get("complaints"), space_before=5)


def add_examination_section(doc, appointment):
    """Компактный абзац анамнеза и объективных данных без лишних подзаголовков."""
    pressure = "—"
    if appointment.get("systolic_pressure") or appointment.get("diastolic_pressure"):
        pressure = (
            f"{clean_value(appointment.get('systolic_pressure'))}/"
            f"{clean_value(appointment.get('diastolic_pressure'))} мм рт. ст."
        )

    bp_note = (appointment.get("bp_note") or "").strip()
    if bp_note and pressure != "—":
        pressure = f"{pressure} ({bp_note})"

    anamnesis_parts = [
        _sentence(appointment.get("life_anamnesis")),
        _sentence(appointment.get("disease_anamnesis")),
    ]

    heredity_description = (appointment.get("heredity_description") or "").strip()
    if heredity_description:
        anamnesis_parts.append(_sentence(heredity_description))
    elif appointment.get("heredity"):
        anamnesis_parts.append("Отягощённая наследственность.")

    comorbidities = (appointment.get("comorbidities") or "").strip()
    if comorbidities:
        anamnesis_parts.append(_sentence(f"Сопутствующие заболевания: {comorbidities}"))

    skin_condition = (appointment.get("skin_condition") or "").strip()
    if skin_condition:
        anamnesis_parts.append(_sentence(f"Кожные покровы: {skin_condition}"))

    anamnesis_parts.extend(
        [
            _sentence(f"Отёки: {clean_value(appointment.get('edema_location'))}"),
            _sentence(f"Артериальное давление: {pressure}"),
            _sentence(f"ЧСС: {clean_value(appointment.get('heart_rate'))}"),
            _sentence(f"Рост: {clean_value(appointment.get('height'))}"),
            _sentence(f"Вес: {clean_value(appointment.get('weight'))}"),
            _sentence(f"ИМТ: {clean_value(appointment.get('bmi'))}"),
        ]
    )

    add_field_inline(doc, "Анамнез", _join_sentences(anamnesis_parts))


def add_lab_sections(doc, context):
    labs = context["labs"]

    add_history_table(
        doc,
        "Общий анализ крови",
        labs["cbc_history"],
        [
            ("Гемоглобин", "hemoglobin"),
            ("Эритроциты", "erythrocytes"),
            ("Лейкоциты", "leukocytes"),
            ("Тромбоциты", "platelets"),
            ("СОЭ", "esr"),
            ("MCV", "mcv"),
            ("Гематокрит", "hematocrit"),
        ],
    )
    add_history_table(
        doc,
        "Биохимический анализ крови",
        labs["biochemistry_history"],
        [
            ("Креатинин", "creatinine"),
            ("Мочевина", "urea"),
            ("Мочевая кислота", "uric_acid"),
            ("Глюкоза", "glucose"),
            ("Общий белок", "total_protein"),
            ("Альбумин", "albumin"),
            ("Калий", "potassium"),
            ("Кальций", "calcium"),
            ("Фосфор", "phosphorus"),
            ("Ферритин", "ferritin"),
            ("ПТГ", "ptg"),
        ],
    )
    add_history_table(
        doc,
        "Общий анализ мочи",
        labs["urinalysis_history"],
        [
            ("Удельный вес", "specific_gravity"),
            ("Белок", "protein"),
            ("Лейкоциты", "leukocytes"),
            ("Эритроциты", "erythrocytes"),
            ("Бактерии", "bacteria"),
        ],
    )
    add_history_table(
        doc,
        "Скорость клубочковой фильтрации",
        labs["metrics_history"],
        [
            ("Креатинин крови, мкмоль/л", "creatinine"),
            ("СКФ CKD-EPI, мл/мин/1.73 м²", "egfr_ckdepi"),
            ("СКФ Кокрофт-Голт, мл/мин", "crcl_cockcroft_gault"),
            ("Категория СКФ", "ckd_stage"),
        ],
    )
    add_history_table(
        doc,
        "Альбуминурия по KDIGO",
        labs["albuminuria_history"],
        [
            ("Альбумин мочи", "urine_albumin_display"),
            ("Креатинин мочи", "urine_creatinine_display"),
            ("ACR, мг/ммоль", "albumin_creatinine_ratio"),
            ("Категория альбуминурии", "albuminuria_category"),
        ],
    )
    add_history_table(
        doc,
        "УЗИ почек",
        labs["ultrasound_history"],
        [
            ("Размер левой почки", "left_kidney_size"),
            ("Размер правой почки", "right_kidney_size"),
            ("Паренхима слева", "left_parenchyma"),
            ("Паренхима справа", "right_parenchyma"),
            ("Описание", "description"),
        ],
    )


def add_conclusion_section(doc, context):
    """Блок заключения: диагнозы МКБ-10 одной строкой и сохранённый прогноз KDIGO."""
    diagnosis_text = icd10_diagnoses_conclusion_text(context.get("diagnoses") or [])
    prognosis_text = kdigo_conclusion_text((context.get("kdigo") or {}).get("current"))

    if not diagnosis_text and not prognosis_text:
        return

    add_table_title(doc, "Заключение")

    if diagnosis_text:
        add_field_inline(doc, "Диагноз", diagnosis_text, space_before=0)

    if prognosis_text:
        add_field_inline(doc, "Прогноз по KDIGO", prognosis_text)


def add_treatment_section(doc, context):
    appointment = context["appointment"]
    medications = context["medications"]
    diet_info = context["diet_info"]

    if medications:
        med_rows = []
        for index, med in enumerate(medications, start=1):
            med_rows.append(
                [
                    index,
                    med.get("medication"),
                    med.get("dosage"),
                    med.get("schedule"),
                ]
            )
        add_small_table(
            doc,
            "Назначения",
            ["№", "Препарат", "Дозировка", "Схема"],
            med_rows,
        )
    else:
        add_field_inline(doc, "Назначения", "—", space_before=5)

    diet = None
    next_control_date = None
    recommendations = None

    if diet_info:
        diet = diet_info.get("diet")
        next_control_date = diet_info.get("next_control_date")
        recommendations = diet_info.get("recommendations")

    if not diet:
        diet = appointment.get("diet")
    if not next_control_date:
        next_control_date = appointment.get("next_control_date")
    if not recommendations:
        recommendations = appointment.get("recommendations")

    recommendations = clean_word_recommendations(recommendations)

    add_field_inline(doc, "Диета", diet, space_before=3)
    if recommendations:
        add_field_inline(doc, "Рекомендации", recommendations)
    add_field_inline(doc, "Дата следующего контроля", fmt_date(next_control_date))


def add_signature_section(doc, appointment):
    doc.add_paragraph()
    signature_paragraph = doc.add_paragraph()
    signature_paragraph.paragraph_format.space_before = Pt(10)
    signature_paragraph.paragraph_format.space_after = Pt(0)
    signature_paragraph.paragraph_format.line_spacing = 1

    appointment_date_text = fmt_date(appointment.get("appointment_date"))
    doctor_name = clean_value(appointment.get("doctor_name"))
    run = signature_paragraph.add_run(
        f"Дата приёма: {appointment_date_text} __________________ / {doctor_name} /"
    )
    set_run_font(run, size=12, bold=False)
