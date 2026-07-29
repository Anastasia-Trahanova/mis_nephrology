"""
Техническое форматирование Word-заключения.

Пустые значения не выводятся: в документе нет прочерков и служебных
заглушек. Обычный текст после строки с ФИО выравнивается по ширине и
получает отступ первой строки 1,25 см.
"""

from __future__ import annotations

import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def has_value(value) -> bool:
    """Возвращает True для реально заполненного значения, включая число 0."""
    if value is None:
        return False
    if isinstance(value, str):
        return bool(value.strip()) and value.strip() != "—"
    return True


def fmt_date(value, with_time: bool = False) -> str:
    """Форматирует date/datetime; пустая дата остаётся пустой."""
    if not value:
        return ""
    try:
        if with_time:
            return value.strftime("%d.%m.%Y %H:%M")
        return value.strftime("%d.%m.%Y")
    except Exception:
        return str(value).strip()


def clean_value(value) -> str:
    """Возвращает пустую строку вместо прочерка для незаполненного поля."""
    if not has_value(value):
        return ""
    return str(value).strip()


def safe_filename(text) -> str:
    """Безопасное имя файла для скачивания."""
    text = text or "patient"
    text = re.sub(r'[\\/*?:"<>|]+', "_", text)
    text = re.sub(r"\s+", "_", text)
    return text.strip("_")[:100]


def set_run_font(run, size: int = 12, bold: bool = False):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def _add_body_paragraph(doc, *, space_before: int = 0, space_after: int = 1):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1
    return paragraph


def add_centered_paragraph(
    doc,
    text,
    size: int = 9,
    bold: bool = False,
    space_after: int = 0,
):
    if not has_value(text):
        return None
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(str(text).strip())
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_field_inline(
    doc,
    title,
    value,
    size: int = 12,
    space_before: int = 0,
    space_after: int = 1,
    title_bold: bool = True,
):
    """Добавляет только заполненное поле: «Название: значение»."""
    if not has_value(value):
        return None

    paragraph = _add_body_paragraph(
        doc,
        space_before=space_before,
        space_after=space_after,
    )

    title_run = paragraph.add_run(f"{title}: ")
    set_run_font(title_run, size=size, bold=title_bold)

    value_run = paragraph.add_run(str(value).strip())
    set_run_font(value_run, size=size, bold=False)
    return paragraph


def add_fields_inline(
    doc,
    fields,
    *,
    size: int = 12,
    separator: str = "; ",
    space_before: int = 0,
    space_after: int = 1,
    title_bold: bool = False,
):
    """Объединяет несколько заполненных полей в один абзац."""
    visible = [(title, value) for title, value in fields if has_value(value)]
    if not visible:
        return None

    paragraph = _add_body_paragraph(
        doc,
        space_before=space_before,
        space_after=space_after,
    )

    for index, (title, value) in enumerate(visible):
        if index:
            separator_run = paragraph.add_run(separator)
            set_run_font(separator_run, size=size, bold=False)

        title_run = paragraph.add_run(f"{title}: ")
        set_run_font(title_run, size=size, bold=title_bold)

        value_run = paragraph.add_run(str(value).strip())
        set_run_font(value_run, size=size, bold=False)

    return paragraph


def add_table_title(doc, title):
    if not has_value(title):
        return None
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(str(title).strip())
    set_run_font(run, size=12, bold=True)
    return paragraph


def format_table_cell(cell, value, bold: bool = False, size: int = 10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(clean_value(value))
    set_run_font(run, size=size, bold=bold)


def add_small_table(doc, title, headers, rows):
    """Компактная таблица с названием раздела."""
    rows = list(rows or [])
    if not rows:
        return None

    add_table_title(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        format_table_cell(header_cells[index], header, bold=True, size=10)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            format_table_cell(cells[index], value, bold=False, size=10)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1

    doc.add_paragraph()
    return table


def add_history_table(doc, title, records, fields, date_key: str = "investigation_date"):
    """
    Добавляет таблицу истории только при наличии данных.

    Пустые показатели не создают строки, а пустые исследования не создают
    столбцы. Первый столбец — показатель, остальные — даты исследований.
    """
    records = list(records or [])
    if not records:
        return None

    visible_fields = [
        (label, key)
        for label, key in fields
        if any(has_value(record.get(key)) for record in records)
    ]
    if not visible_fields:
        return None

    visible_records = [
        record
        for record in records
        if any(has_value(record.get(key)) for _, key in visible_fields)
    ]
    if not visible_records:
        return None

    add_table_title(doc, title)

    headers = ["Показатель"] + [
        fmt_date(record.get(date_key)) for record in visible_records
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        format_table_cell(header_cells[index], header, bold=True, size=10)

    for label, key in visible_fields:
        cells = table.add_row().cells
        format_table_cell(cells[0], label, bold=True, size=10)
        for index, record in enumerate(visible_records, start=1):
            format_table_cell(cells[index], record.get(key), bold=False, size=10)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1

    doc.add_paragraph()
    return table


def unit_label(value) -> str:
    """Человекочитаемые единицы измерения для Word-отчёта."""
    labels = {
        "mg_l": "мг/л",
        "g_l": "г/л",
        "mmol_l": "ммоль/л",
        "umol_l": "мкмоль/л",
    }
    return labels.get(value, value or "")


def value_with_unit(value, unit) -> str:
    if not has_value(value):
        return ""
    unit_text = unit_label(unit)
    return f"{value} {unit_text}".strip()
