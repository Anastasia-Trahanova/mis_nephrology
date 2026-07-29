"""Подпись врача в Word-заключении."""

from __future__ import annotations

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .formatting import clean_value, fmt_date, set_run_font


def _remove_table_borders(table):
    """Убирает видимые границы служебной таблицы подписи."""
    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)

    for edge_name in (
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ):
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))

        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)

        edge.set(qn("w:val"), "nil")


def _format_signature_paragraph(paragraph, alignment):
    paragraph.alignment = alignment
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1


def add_signature_section(doc, appointment):
    doc.add_paragraph()

    appointment_date = fmt_date(appointment.get("appointment_date"))
    doctor_name = clean_value(appointment.get("doctor_name"))

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _remove_table_borders(table)

    left_cell, right_cell = table.rows[0].cells
    left_cell.width = Cm(7)
    right_cell.width = Cm(11)

    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left_paragraph = left_cell.paragraphs[0]
    _format_signature_paragraph(
        left_paragraph,
        WD_ALIGN_PARAGRAPH.LEFT,
    )

    if appointment_date:
        left_run = left_paragraph.add_run(
            f"Дата приёма: {appointment_date}"
        )
        set_run_font(left_run, size=12, bold=False)

    right_paragraph = right_cell.paragraphs[0]
    _format_signature_paragraph(
        right_paragraph,
        WD_ALIGN_PARAGRAPH.RIGHT,
    )

    signature_text = "__________________"
    if doctor_name:
        signature_text += f" / {doctor_name} /"

    right_run = right_paragraph.add_run(signature_text)
    set_run_font(right_run, size=12, bold=False)