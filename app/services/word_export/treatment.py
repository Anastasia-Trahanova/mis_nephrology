"""Диета, рекомендации, медикаментозная терапия и дата контроля для Word."""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.medication_therapy import build_medication_therapy_groups

from .formatting import add_field_inline, fmt_date, has_value, set_run_font
from .text import clean_word_recommendations


def _add_medication_paragraph(doc, medication) -> None:
    """Добавляет препарат так же, как в карточке: жирное название и обычные детали."""
    name = str(medication.get("medication") or "").strip()
    if not name:
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(1.25)
    paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1

    bullet_run = paragraph.add_run("• ")
    set_run_font(bullet_run, size=12, bold=False)

    name_run = paragraph.add_run(name)
    set_run_font(name_run, size=12, bold=True)

    if has_value(medication.get("dosage")):
        dosage_run = paragraph.add_run(
            f" — {str(medication.get('dosage')).strip()}"
        )
        set_run_font(dosage_run, size=12, bold=False)

    if has_value(medication.get("schedule")):
        schedule_run = paragraph.add_run(
            f" ({str(medication.get('schedule')).strip()})"
        )
        set_run_font(schedule_run, size=12, bold=False)


def _add_medications_list(doc, medications) -> None:
    groups = build_medication_therapy_groups(prescriptions=medications)
    visible_groups = [
        group for group in groups if group.get("prescriptions")
    ]
    if not visible_groups:
        add_field_inline(doc, "Медикаментозная терапия", "—")
        return

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.left_indent = Cm(0)
    title_paragraph.paragraph_format.first_line_indent = Cm(1.25)
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(1)
    title_paragraph.paragraph_format.line_spacing = 1
    title_run = title_paragraph.add_run("Медикаментозная терапия:")
    set_run_font(title_run, size=12, bold=True)

    for group in visible_groups:
        group_paragraph = doc.add_paragraph()
        group_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        group_paragraph.paragraph_format.left_indent = Cm(0)
        group_paragraph.paragraph_format.first_line_indent = Cm(0)
        group_paragraph.paragraph_format.space_before = Pt(1)
        group_paragraph.paragraph_format.space_after = Pt(0)
        group_paragraph.paragraph_format.line_spacing = 1
        group_run = group_paragraph.add_run(f'{group["title"]}:')
        set_run_font(group_run, size=12, bold=False)

        for medication in group.get("prescriptions") or []:
            _add_medication_paragraph(doc, medication)


def add_treatment_section(doc, context):
    appointment = context["appointment"]
    medications = context.get("medications") or []
    diet_info = context.get("diet_info") or {}

    diet = diet_info.get("diet") or appointment.get("diet")
    recommendations = (
        diet_info.get("recommendations")
        or appointment.get("recommendations")
    )
    recommendations = clean_word_recommendations(recommendations)
    next_control_date = (
        diet_info.get("next_control_date")
        or appointment.get("next_control_date")
    )

    add_field_inline(doc, "Диета", diet, space_before=0)
    add_field_inline(doc, "Рекомендации", recommendations)
    _add_medications_list(doc, medications)

    next_date_paragraph = add_field_inline(
        doc,
        "Дата следующего контроля",
        fmt_date(next_control_date),
    )
    if next_date_paragraph is not None:
        next_date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        next_date_paragraph.paragraph_format.left_indent = Cm(0)
        next_date_paragraph.paragraph_format.first_line_indent = Cm(1.25)
