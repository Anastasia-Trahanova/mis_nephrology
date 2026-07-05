from fastapi import APIRouter, Request, HTTPException
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, StreamingResponse
from io import BytesIO
from urllib.parse import quote
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from datetime import datetime
from datetime import date
from app.repositories.patients import (
    get_all_patients,
    get_patient_by_id,
)
from app.repositories.appointments import (
    get_all_appointments,
    get_patient_appointments,
    get_appointment_full_data,
    get_last_appointment_data,
    get_appointment_medications,
    get_appointment_diet,
)
from app.repositories.reference_data import (
    get_branches,
    get_locations_by_branch,
    get_doctors,
    get_location_info,
)
from app.repositories.lab_history import (
    get_patient_biochemistry_history,
    get_patient_cbc_history,
    get_patient_urinalysis_history,
    get_patient_ultrasound_history,
    get_patient_metrics_history,
    get_patient_albuminuria_history,
)
from app.repositories.ckd_prognosis import (
    get_patient_ckd_prognosis_history,
    get_appointment_ckd_prognosis,
)
from app.services.patient_card_context_service import get_patient_card_context
from app.services.appointment_form_context_service import (
    get_new_appointment_context,
    get_new_patient_context,
)
router = APIRouter(tags=["pages"]) #	Создаёт группу маршрутов с тегом pages (для документации)
templates = Jinja2Templates(directory="app/templates") #	Указывает, где искать HTML-файлы


@router.get("/", response_class=HTMLResponse) #Обрабатывает GET-запросы на главную страницу? Возвращается HTML (не JSON)
def home(request: Request): #Асинхронная функция
    branches = get_branches() #	Загружает список филиалов для фильтра
    doctors = get_doctors()
    return templates.TemplateResponse("index.html", { #Рендерит HTML-шаблон index.html и передаёт в него данные
        "request": request,
        "branches": branches,
        "doctors": doctors
    })


@router.get("/patients", response_class=HTMLResponse)
def patients_list(request: Request, search: str = None, limit: int = 500, offset: int = 0):
    """Страница списка пациентов."""
    patients = get_all_patients(search=search, limit=limit, offset=offset)
    return templates.TemplateResponse("patients_list.html", {
        "request": request,
        "patients": patients,
        "search": search or "",
        "limit": limit,
        "offset": offset,
    })




@router.get("/patient/{patient_id}", response_class=HTMLResponse)
def patient_card(request: Request, patient_id: int):
    """Карточка пациента. Оптимизировано: данные собираются одним соединением к БД."""
    show_form = request.query_params.get('show_form') == 'true'
    show_previous_labs = request.query_params.get('show_previous_labs') == 'true'

    selected_appointment_id = request.query_params.get('appointment_id')
    if selected_appointment_id:
        try:
            selected_appointment_id = int(selected_appointment_id)
        except ValueError:
            raise HTTPException(status_code=400, detail="Некорректный appointment_id")
    else:
        selected_appointment_id = None

    context = get_patient_card_context(
        patient_id=patient_id,
        selected_appointment_id=selected_appointment_id,
        show_previous_labs=show_previous_labs,
        show_form=show_form,
    )

    if not context:
        raise HTTPException(status_code=404, detail="Пациент не найден")

    if context.get("forbidden"):
        raise HTTPException(status_code=403, detail="Этот приём не принадлежит данному пациенту")

    now = datetime.now()
    context.update({
        "request": request,
        "now_date": now.strftime('%Y-%m-%d'),
        "now_time": now.strftime('%H:%M'),
        "show_form": show_form,
    })

    return templates.TemplateResponse("patient_card.html", context)

@router.get("/api/patient/{patient_id}/biochemistry_history")
async def api_biochemistry_history(patient_id: int):
    return get_patient_biochemistry_history(patient_id)

@router.get("/new-patient", response_class=HTMLResponse)
def new_patient_form(request: Request):
    context = get_new_patient_context()

    now = datetime.now()
    context.update({
        "request": request,
        "now_date": now.strftime("%Y-%m-%d"),
        "now_time": now.strftime("%H:%M"),
    })

    return templates.TemplateResponse("new_patient.html", context)

@router.get("/new-appointment/{patient_id}", response_class=HTMLResponse)
def new_appointment_form(request: Request, patient_id: int):
    """Форма нового приёма. Оптимизировано: данные собираются одним соединением к БД."""
    context = get_new_appointment_context(patient_id)
    if not context:
        raise HTTPException(status_code=404, detail="Пациент не найден")

    now = datetime.now()
    context.update({
        "request": request,
        "now_date": now.strftime("%Y-%m-%d"),
        "now_time": now.strftime("%H:%M"),
    })

    return templates.TemplateResponse("new_appointment.html", context)


@router.get("/export/{appointment_id}")
def export_appointment(appointment_id: int):
    """Экспорт заключения в текстовом виде"""
    appointment = get_appointment_full_data(appointment_id)
    if not appointment:
        raise HTTPException(status_code=404, detail="Приём не найден")
    
    # Получаем информацию об отделении (как в patient_card)
    location_info = None
    header_text = ""
    
    if appointment.get('location_id'):
        location_info = get_location_info(appointment['location_id'])
        
        # Формируем шапку на основе данных из БД
        if location_info:
            company_name = location_info.get('company_name') or 'ООО «КОМПАНИЯ «ФЕСФАРМ»'
            location_name = location_info.get('location_name') or ''
            location_address = location_info.get('location_address') or ''
            branch_phone = location_info.get('branch_phone') or ''
            company_phone = location_info.get('company_phone') or ''
            phone = branch_phone or company_phone or ''
            branch_email = location_info.get('branch_email') or ''
            company_email = location_info.get('company_email') or ''
            email = branch_email or company_email or ''
            
            header_text = f"""
{company_name}
{location_name}
{location_address}
Тел: {phone}{' | Email: ' + email if email else ''}

"""
        else:
            header_text = "ООО «КОМПАНИЯ «ФЕСФАРМ»\n\n"
    else:
        header_text = "ООО «КОМПАНИЯ «ФЕСФАРМ»\n\n"
    
    # Формируем текст заключения
    report = f"""{header_text}
ЗАКЛЮЧЕНИЕ ПО РЕЗУЛЬТАТАМ ПРИЁМА

Дата приёма: {appointment['appointment_date'].strftime('%d.%m.%Y %H:%M') if appointment['appointment_date'] else '—'}
Врач: {appointment['doctor_name']}
Отделение: {appointment.get('location_name', '—')} ({appointment.get('branch_name', '—')})

Пациент: {appointment.get('patient_last_name', '')} {appointment.get('patient_first_name', '')} {appointment.get('patient_patronymic', '')}

Жалобы: {appointment.get('complaints', '—')}

Диагноз: {appointment.get('main_diagnosis', '—')}

Лечение: {appointment.get('medication', '—')} {appointment.get('dosage', '')} {appointment.get('schedule', '')}

Дата следующего контроля: {appointment['next_control_date'].strftime('%d.%m.%Y') if appointment.get('next_control_date') else '—'}
"""
    
    return HTMLResponse(content=f"<pre>{report}</pre>", media_type="text/html; charset=utf-8")

def fmt_date(value, with_time=False):
    """Форматирует date/datetime для заключения."""
    if not value:
        return "—"
    try:
        if with_time:
            return value.strftime("%d.%m.%Y %H:%M")
        return value.strftime("%d.%m.%Y")
    except Exception:
        return str(value)


def clean_value(value):
    """Красиво выводит пустые значения."""
    if value is None or value == "":
        return "—"
    return str(value)


def safe_filename(text):
    """Безопасное имя файла для скачивания."""
    text = text or "patient"
    text = re.sub(r'[\\/*?:"<>|]+', "_", text)
    text = re.sub(r"\s+", "_", text)
    return text.strip("_")[:100]


def set_run_font(run, size=12, bold=False):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_centered_paragraph(doc, text, size=9, bold=False, space_after=0):
    if not text:
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1

    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold)


def add_field_inline(doc, title, value, size=12, space_before=0, space_after=1):
    """
    Одно поле = одна строка/один абзац:
    Жалобы: текст жалоб
    """
    if value is None or value == "":
        value = "—"

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1

    title_run = paragraph.add_run(f"{title}: ")
    set_run_font(title_run, size=size, bold=True)

    value_run = paragraph.add_run(str(value))
    set_run_font(value_run, size=size, bold=False)


def add_table_title(doc, title):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1

    run = paragraph.add_run(title)
    set_run_font(run, size=12, bold=True)


def format_table_cell(cell, value, bold=False, size=10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    run = paragraph.add_run(clean_value(value))
    set_run_font(run, size=size, bold=bold)


def add_small_table(doc, title, headers, rows):
    """
    Компактная таблица с названием исследования.
    """
    add_table_title(doc, title)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        format_table_cell(header_cells[index], header, bold=True, size=10)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            format_table_cell(cells[index], value, bold=False, size=10)

    # Уменьшаем интервалы внутри всех ячеек
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1

    doc.add_paragraph()

def add_history_table(doc, title, records, fields, date_key="investigation_date"):
    """
    Таблица истории анализов/расчётов:
    первый столбец — показатель,
    дальше столбцы по датам исследования/оценки.
    date_key позволяет использовать не только investigation_date,
    но и assessment_date для прогноза ХБП.
    """
    add_table_title(doc, title)

    if not records:
        add_field_inline(doc, title, "нет данных")
        return

    headers = ["Показатель"]
    for record in records:
        headers.append(fmt_date(record.get(date_key)))

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    # Заголовки
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        format_table_cell(header_cells[index], header, bold=True, size=10)

    # Строки показателей
    for label, key in fields:
        cells = table.add_row().cells
        format_table_cell(cells[0], label, bold=True, size=10)

        for index, record in enumerate(records, start=1):
            format_table_cell(cells[index], record.get(key), bold=False, size=10)

    # Уплотняем строки
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1

    doc.add_paragraph()


def unit_label(value):
    """Человекочитаемые единицы измерения для Word-отчёта."""
    labels = {
        "mg_l": "мг/л",
        "g_l": "г/л",
        "mmol_l": "ммоль/л",
        "umol_l": "мкмоль/л",
    }
    return labels.get(value, value or "")


def value_with_unit(value, unit):
    if value is None or value == "":
        return "—"
    unit_text = unit_label(unit)
    return f"{value} {unit_text}".strip()


def prepare_albuminuria_records(records):
    """Добавляет поля для красивого вывода альбуминурии в Word."""
    prepared = []
    for record in records or []:
        item = dict(record)
        item["urine_albumin_display"] = value_with_unit(
            item.get("urine_albumin"),
            item.get("urine_albumin_unit"),
        )
        item["urine_creatinine_display"] = value_with_unit(
            item.get("urine_creatinine"),
            item.get("urine_creatinine_unit"),
        )
        prepared.append(item)
    return prepared


def prognosis_display(record):
    if not record:
        return "—"

    combined = clean_value(record.get("combined_category"))
    text = clean_value(record.get("prognosis_text"))

    if combined == "—" and text == "—":
        return "—"

    if combined == "—":
        return text

    if text == "—":
        return combined

    return f"{combined}: {text}"


def prepare_ckd_prognosis_records(records):
    """Добавляет поле prognosis_display для компактного вывода прогноза ХБП."""
    prepared = []
    for record in records or []:
        item = dict(record)
        item["prognosis_display"] = prognosis_display(item)
        prepared.append(item)
    return prepared


@router.get("/export/{appointment_id}/docx")
def export_appointment_docx(appointment_id: int):
    """Экспорт заключения приёма в настоящий Word-файл .docx"""

    appointment = get_appointment_full_data(appointment_id)
    if not appointment:
        raise HTTPException(status_code=404, detail="Приём не найден")
 
    
    medications = get_appointment_medications(appointment_id)
    diet_info = get_appointment_diet(appointment_id)

    patient_id = appointment.get("patient_id")

    until_date = None
    if appointment.get("appointment_date"):
        until_date = appointment["appointment_date"].date()

    cbc_history = get_patient_cbc_history(patient_id, until_date)
    biochemistry_history = get_patient_biochemistry_history(patient_id, until_date)
    urinalysis_history = get_patient_urinalysis_history(patient_id, until_date)
    metrics_history = get_patient_metrics_history(patient_id, until_date)
    albuminuria_history = prepare_albuminuria_records(
        get_patient_albuminuria_history(patient_id, until_date)
    )
    ckd_prognosis_history = prepare_ckd_prognosis_records(
        get_patient_ckd_prognosis_history(patient_id, until_date)
    )
    ckd_prognosis_current = get_appointment_ckd_prognosis(appointment_id)
    ultrasound_history = get_patient_ultrasound_history(patient_id, until_date)

    location_info = None
    if appointment.get("location_id"):
        location_info = get_location_info(appointment["location_id"])

    doc = Document()

    # Настройки страницы
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.3)

    # Основной стиль
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # =========================
    # ШАПКА КЛИНИКИ В 2 СТРОКИ
    # =========================
    if location_info:
        company_name = location_info.get("company_name") or "ООО «КОМПАНИЯ «ФЕСФАРМ»"
        location_name = location_info.get("location_name") or ""
        location_address = location_info.get("location_address") or ""

        branch_phone = location_info.get("branch_phone") or ""
        company_phone = location_info.get("company_phone") or ""
        phone = branch_phone or company_phone or ""

        branch_email = location_info.get("branch_email") or ""
        company_email = location_info.get("company_email") or ""
        email = branch_email or company_email or ""

        first_line = company_name
        if location_name:
            first_line += f" — {location_name}"

        second_line_parts = []

        if location_address:
            second_line_parts.append(location_address)

        if phone:
            second_line_parts.append(f"Тел: {phone}")

        if email:
            second_line_parts.append(f"Email: {email}")

        add_centered_paragraph(doc, first_line, size=9, bold=False, space_after=0)
        add_centered_paragraph(doc, " | ".join(second_line_parts), size=9, bold=False, space_after=4)
    else:
        add_centered_paragraph(doc, "ООО «КОМПАНИЯ «ФЕСФАРМ»", size=9, bold=False, space_after=4)

    # =========================
    # ЗАГОЛОВОК
    # =========================
    add_centered_paragraph(doc, "КОНСУЛЬТАТИВНОЕ ЗАКЛЮЧЕНИЕ", size=14, bold=True, space_after=6)

    # Убираем верхний служебный блок:
    # Дата приёма / Врач / Отделение здесь больше НЕ выводятся

    # =========================
    # ПАЦИЕНТ
    # =========================
    add_field_inline(doc, "Пациент", appointment.get("patient_fio"))
    add_field_inline(doc, "Дата рождения", fmt_date(appointment.get("birth_date")))

    if appointment.get("age_at_appointment") is not None:
        try:
            age_text = f"{int(appointment.get('age_at_appointment'))} лет"
        except Exception:
            age_text = str(appointment.get("age_at_appointment"))

        add_field_inline(doc, "Возраст на момент приёма", age_text)

    # =========================
    # ОПРОС — заголовок не пишем
    # =========================
    add_field_inline(doc, "Жалобы", appointment.get("complaints"), space_before=5)
    add_field_inline(doc, "Анамнез жизни", appointment.get("life_anamnesis"))
    add_field_inline(doc, "Анамнез заболевания", appointment.get("disease_anamnesis"))

    heredity = "Да" if appointment.get("heredity") else "Нет"
    add_field_inline(doc, "Отягощённая наследственность", heredity)
    add_field_inline(doc, "Описание наследственности", appointment.get("heredity_description"))
    add_field_inline(doc, "Сопутствующие заболевания", appointment.get("comorbidities"))

    # =========================
    # ОСМОТР — заголовок не пишем
    # =========================
    add_field_inline(doc, "Кожные покровы", appointment.get("skin_condition"), space_before=5)
    add_field_inline(doc, "Отёки", appointment.get("edema_location"))

    pressure = "—"
    if appointment.get("systolic_pressure") or appointment.get("diastolic_pressure"):
        pressure = f"{clean_value(appointment.get('systolic_pressure'))}/{clean_value(appointment.get('diastolic_pressure'))} мм рт. ст."

    add_field_inline(doc, "Артериальное давление", pressure)
    add_field_inline(doc, "Примечание к АД", appointment.get("bp_note"))
    add_field_inline(doc, "ЧСС", appointment.get("heart_rate"))
    add_field_inline(doc, "Рост", appointment.get("height"))
    add_field_inline(doc, "Вес", appointment.get("weight"))
    add_field_inline(doc, "ИМТ", appointment.get("bmi"))

    # =========================
    # ЛАБОРАТОРНЫЕ ИССЛЕДОВАНИЯ
    # Общий заголовок не пишем.
    # Каждая таблица имеет своё название.
    # =========================
    
    add_history_table(
    doc,
    "Общий анализ крови",
    cbc_history,
    [
        ("Гемоглобин", "hemoglobin"),
        ("Эритроциты", "erythrocytes"),
        ("Лейкоциты", "leukocytes"),
        ("Тромбоциты", "platelets"),
        ("СОЭ", "esr"),
        ("MCV", "mcv"),
        ("Гематокрит", "hematocrit"),
    ],
    )

    add_history_table(
        doc,
        "Биохимический анализ крови",
        biochemistry_history,
        [
            ("Креатинин", "creatinine"),
            ("Мочевина", "urea"),
            ("Мочевая кислота", "uric_acid"),
            ("Глюкоза", "glucose"),
            ("Общий белок", "total_protein"),
            ("Альбумин", "albumin"),
            ("Калий", "potassium"),
            ("Кальций", "calcium"),
            ("Фосфор", "phosphorus"),
            ("Ферритин", "ferritin"),
            ("ПТГ", "ptg"),
        ],
    )

    add_history_table(
        doc,
        "Общий анализ мочи",
        urinalysis_history,
        [
            ("Удельный вес", "specific_gravity"),
            ("Белок", "protein"),
            ("Лейкоциты", "leukocytes"),
            ("Эритроциты", "erythrocytes"),
            ("Бактерии", "bacteria"),
        ],
    )

    # =========================
    # РАСЧЁТНЫЕ ПОКАЗАТЕЛИ
    # После ОАМ, но без отдельного заголовка
    # =========================
    add_history_table(
        doc,
        "Скорость клубочковой фильтрации",
        metrics_history,
        [
            ("Креатинин крови, мкмоль/л", "creatinine"),
            ("СКФ CKD-EPI, мл/мин/1.73 м²", "egfr_ckdepi"),
            ("СКФ Кокрофт-Голт, мл/мин", "crcl_cockcroft_gault"),
            ("Категория СКФ", "ckd_stage"),
        ],
    )

    add_history_table(
        doc,
        "Альбуминурия по KDIGO",
        albuminuria_history,
        [
            ("Альбумин мочи", "urine_albumin_display"),
            ("Креатинин мочи", "urine_creatinine_display"),
            ("ACR, мг/ммоль", "albumin_creatinine_ratio"),
            ("Категория альбуминурии", "albuminuria_category"),
        ],
    )

    add_history_table(
        doc,
        "Прогноз ХБП по KDIGO",
        ckd_prognosis_history,
        [
            ("Категория СКФ", "gfr_category"),
            ("Категория альбуминурии", "albuminuria_category"),
            ("Итоговая категория", "combined_category"),
            ("Прогноз", "prognosis_text"),
        ],
        date_key="assessment_date",
    )

    if ckd_prognosis_current:
        add_field_inline(
            doc,
            "Актуальный прогноз ХБП",
            prognosis_display(ckd_prognosis_current),
            space_before=2,
        )

    # =========================
    # УЗИ после расчётных показателей
    # =========================
    add_history_table(
    doc,
    "УЗИ почек",
    ultrasound_history,
    [
        ("Размер левой почки", "left_kidney_size"),
        ("Размер правой почки", "right_kidney_size"),
        ("Паренхима слева", "left_parenchyma"),
        ("Паренхима справа", "right_parenchyma"),
        ("Описание", "description"),
    ],
    )

    # =========================
    # ДИАГНОЗЫ — заголовок не пишем
    # =========================
    add_field_inline(doc, "Основной диагноз", appointment.get("main_diagnosis"), space_before=5)
    add_field_inline(doc, "Осложнения", appointment.get("complications"))
    add_field_inline(doc, "Сопутствующие диагнозы", appointment.get("diag_comorbidities"))

    # =========================
    # ЛЕЧЕНИЕ
    # =========================
    if medications:
        med_rows = []
        for index, med in enumerate(medications, start=1):
            med_rows.append([
                index,
                med.get("medication"),
                med.get("dosage"),
                med.get("schedule"),
            ])

        add_small_table(
            doc,
            "Назначения",
            ["№", "Препарат", "Дозировка", "Схема"],
            med_rows,
        )
    else:
        add_field_inline(doc, "Назначения", "—", space_before=5)

    # Диета, рекомендации и контроль
    diet = None
    next_control_date = None
    recommendations = None

    if diet_info:
        diet = diet_info.get("diet")
        next_control_date = diet_info.get("next_control_date")
        recommendations = diet_info.get("recommendations")

    if not diet:
        diet = appointment.get("diet")

    if not next_control_date:
        next_control_date = appointment.get("next_control_date")

    if not recommendations:
        recommendations = appointment.get("recommendations")

    add_field_inline(doc, "Диета", diet, space_before=3)
    add_field_inline(doc, "Рекомендации", recommendations)
    add_field_inline(doc, "Дата следующего контроля", fmt_date(next_control_date))

    # =========================
    # ПОДПИСЬ ВРАЧА
    # =========================
    doc.add_paragraph()

    signature_paragraph = doc.add_paragraph()
    signature_paragraph.paragraph_format.space_before = Pt(10)
    signature_paragraph.paragraph_format.space_after = Pt(0)
    signature_paragraph.paragraph_format.line_spacing = 1

    appointment_date_text = fmt_date(appointment.get("appointment_date"))
    doctor_name = clean_value(appointment.get("doctor_name"))

    run = signature_paragraph.add_run(
        f"Дата приёма: {appointment_date_text}        __________________ / {doctor_name} /"
    )
    set_run_font(run, size=12, bold=False)

    # Сохраняем в память
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    patient_fio = safe_filename(appointment.get("patient_fio"))
    appointment_date = fmt_date(appointment.get("appointment_date")).replace(".", "-")
    filename = f"Заключение_{patient_fio}_{appointment_date}.docx"

    quoted_filename = quote(filename)

    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{quoted_filename}"
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )