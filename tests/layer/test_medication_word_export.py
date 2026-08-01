"""Тесты данных и форматирования групп лекарств в Word-заключении."""

from __future__ import annotations

from datetime import datetime

from docx import Document

import app.services.word_export.data as word_data
from app.services.word_export.treatment import _add_medications_list


def test_word_medications_are_grouped_in_configured_order():
    document = Document()
    medications = [
        {
            "therapy_group": "Другие препараты",
            "medication": "Аллопуринол",
            "dosage": "100 мг",
            "schedule": "после еды",
        },
        {
            "therapy_group": "Коррекция анемии",
            "medication": "Эпоэтин альфа",
            "dosage": "4000 МЕ",
            "schedule": "3 раза в неделю",
        },
        {
            "therapy_group": "Коррекция АД, ЧСС",
            "medication": "Лозартан",
            "dosage": "50 мг",
            "schedule": "утром",
        },
    ]

    _add_medications_list(document, medications)
    texts = [paragraph.text for paragraph in document.paragraphs]

    assert texts == [
        "Медикаментозная терапия:",
        "Препараты для коррекции АД, ЧСС:",
        "• Лозартан — 50 мг (утром)",
        "Препараты для коррекции анемии:",
        "• Эпоэтин альфа — 4000 МЕ (3 раза в неделю)",
        "Дополнительно:",
        "• Аллопуринол — 100 мг (после еды)",
    ]


def test_word_group_heading_has_no_indent_colon_or_bold_and_drug_name_is_bold():
    document = Document()
    _add_medications_list(
        document,
        [
            {
                "therapy_group": "Нефропротекция",
                "medication": "Финеренон",
                "dosage": "10 мг",
                "schedule": "1 раз в день",
            }
        ],
    )

    group_paragraph = next(p for p in document.paragraphs if p.text == "Нефропротекторные препараты:")
    medication_paragraph = next(p for p in document.paragraphs if "Финеренон" in p.text)

    assert int(group_paragraph.paragraph_format.left_indent or 0) == 0
    assert int(group_paragraph.paragraph_format.first_line_indent or 0) == 0
    assert all(run.bold is False for run in group_paragraph.runs)
    assert medication_paragraph.runs[0].text == "• "
    assert medication_paragraph.runs[0].bold is False
    assert medication_paragraph.runs[1].text == "Финеренон"
    assert medication_paragraph.runs[1].bold is True
    assert all(run.bold is False for run in medication_paragraph.runs[2:])


def test_word_skips_empty_name_and_handles_missing_dose_or_schedule():
    document = Document()
    _add_medications_list(
        document,
        [
            {
                "therapy_group": "Другие препараты",
                "medication": "Аллопуринол",
                "dosage": None,
                "schedule": "после еды",
            },
            {
                "therapy_group": "Другие препараты",
                "medication": "Фебуксостат",
                "dosage": "40 мг",
                "schedule": None,
            },
            {
                "therapy_group": "Другие препараты",
                "medication": "",
                "dosage": "не должно выводиться",
                "schedule": "не должно выводиться",
            },
        ],
    )

    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "• Аллопуринол (после еды)" in texts
    assert "• Фебуксостат — 40 мг" in texts
    assert all("не должно выводиться" not in text for text in texts)


def test_word_context_loads_medications_from_prescriptions_repository(monkeypatch):
    appointment = {
        "id": 13,
        "patient_id": 5,
        "location_id": None,
        "appointment_date": datetime(2026, 4, 5, 10, 0),
    }
    medications = [
        {
            "therapy_group": "Коррекция гиперлипидемии",
            "medication": "Розувастатин",
            "dosage": "10 мг",
            "schedule": "вечером",
        }
    ]
    medication_calls: list[int] = []

    monkeypatch.setattr(word_data, "get_appointment_full_data", lambda appointment_id: appointment)
    monkeypatch.setattr(
        word_data,
        "get_appointment_medications",
        lambda appointment_id: medication_calls.append(appointment_id) or medications,
    )
    monkeypatch.setattr(word_data, "get_appointment_diet", lambda *_: None)
    monkeypatch.setattr(word_data, "get_appointment_icd10_diagnoses", lambda *_: [])
    monkeypatch.setattr(word_data, "get_location_info", lambda *_: None)
    monkeypatch.setattr(word_data, "get_appointment_ckd_prognosis", lambda *_: None)
    monkeypatch.setattr(word_data, "get_patient_ckd_prognosis_history", lambda *_: [])
    for name in (
        "get_patient_albuminuria_history",
        "get_patient_biochemistry_history",
        "get_patient_cbc_history",
        "get_patient_metrics_history",
        "get_patient_ultrasound_history",
        "get_patient_urinalysis_history",
    ):
        monkeypatch.setattr(word_data, name, lambda *_: [])

    context = word_data.get_word_export_context(13)

    assert medication_calls == [13]
    assert context["medications"] == medications
